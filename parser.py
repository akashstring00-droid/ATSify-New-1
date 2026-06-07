import io
from pypdf import PdfReader
from docx import Document

def parse_pdf(file_bytes):
    """
    Extracts text from a PDF file in bytes format.
    """
    text = ""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")
    return text.strip()

def parse_docx(file_bytes):
    """
    Extracts text from a Word document in bytes format, including paragraphs and tables.
    """
    text = []
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        # Read paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        
        # Read tables to catch table-based resumes
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Clean up cell paragraphs
                    cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text.append(" | ".join(row_text))
                    
    except Exception as e:
        raise ValueError(f"Error parsing Word DOCX file: {str(e)}")
    
    return "\n".join(text).strip()

def extract_text(uploaded_file):
    """
    Identifies the file type and extracts the text content.
    """
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    
    # Reset read pointer just in case
    uploaded_file.seek(0)
    
    if file_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif file_name.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX resume.")
