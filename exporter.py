import io
import os
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib import colors

# Load templates configuration fallback
from templates import TEMPLATES

# Helper to convert Hex to RGBColor for docx
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    if len(hex_str) == 3:
        hex_str = ''.join([c*2 for c in hex_str])
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

# Helper to clear table cell margins in docx
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper to calculate approximate word count for adaptive compression
def get_resume_word_count(data):
    wc = 0
    wc += len(data.get("professional_summary", "").split())
    # Skills
    sk = data.get("skills", {})
    wc += len(sk.get("technical_skills", [])) * 2
    wc += len(sk.get("soft_skills", [])) * 2
    wc += len(sk.get("tools_frameworks", [])) * 2
    # Experience
    for job in data.get("work_experience", []):
        wc += len(job.get("job_title", "").split())
        wc += len(job.get("company", "").split())
        wc += sum(len(bp.split()) for bp in job.get("bullet_points", []))
    # Projects
    for proj in data.get("projects", []):
        wc += len(proj.get("name", "").split())
        wc += sum(len(bp.split()) for bp in proj.get("description", []))
    # Education
    for edu in data.get("education", []):
        wc += len(edu.get("degree", "").split())
        wc += len(edu.get("institution", "").split())
    return wc


# ==============================================================================
# DOCX MODULAR BUILDING SYSTEM
# ==============================================================================

def add_section_heading_docx(container, text, font_name, primary_rgb, spacer_before=12, accent_style="underline", primary_color_hex="000000"):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(spacer_before)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text.upper())
    run.font.name = font_name
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = primary_rgb
    
    if accent_style == "underline":
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="3" w:color="E2E8F0"/></w:pBdr>')
        pPr.append(pBdr)
    elif accent_style == "left_border":
        pPr = p._p.get_or_add_pPr()
        clean_color = primary_color_hex.lstrip('#')
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="6" w:color="{clean_color}"/></w:pBdr>')
        pPr.append(pBdr)
        p.paragraph_format.left_indent = Inches(0.1)
    return p

def render_summary_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, accent_style="underline", primary_color_hex="000000"):
    summary = data.get('professional_summary', '')
    if summary:
        add_section_heading_docx(container, "Professional Summary", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        p = container.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(summary)
        run.font.name = font_name
        run.font.size = Pt(body_size)

def render_skills_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, accent_style="underline", primary_color_hex="000000"):
    skills_data = data.get('skills', {})
    if skills_data:
        add_section_heading_docx(container, "Skills", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        tech_skills = skills_data.get('technical_skills', [])
        soft_skills = skills_data.get('soft_skills', [])
        tools = skills_data.get('tools_frameworks', [])
        
        if tech_skills:
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            r_lbl = p.add_run("Technical Skills: ")
            r_lbl.bold = True
            r_lbl.font.name = font_name
            r_lbl.font.size = Pt(body_size)
            r_val = p.add_run(", ".join(tech_skills))
            r_val.font.name = font_name
            r_val.font.size = Pt(body_size)
            
        if soft_skills:
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            r_lbl = p.add_run("Soft Skills: ")
            r_lbl.bold = True
            r_lbl.font.name = font_name
            r_lbl.font.size = Pt(body_size)
            r_val = p.add_run(", ".join(soft_skills))
            r_val.font.name = font_name
            r_val.font.size = Pt(body_size)
            
        if tools:
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.1
            r_lbl = p.add_run("Tools & Systems: ")
            r_lbl.bold = True
            r_lbl.font.name = font_name
            r_lbl.font.size = Pt(body_size)
            r_val = p.add_run(", ".join(tools))
            r_val.font.name = font_name
            r_val.font.size = Pt(body_size)

def render_experience_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, width_col1=4.5, width_col2=2.0, accent_style="underline", primary_color_hex="000000"):
    jobs = data.get('work_experience', [])
    if jobs:
        add_section_heading_docx(container, "Professional Experience", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        for job in jobs:
            table = container.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(width_col1)
            table.columns[1].width = Inches(width_col2)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(1)
            
            run_title = p_left.add_run(job.get('job_title', ''))
            run_title.bold = True
            run_title.font.name = font_name
            run_title.font.size = Pt(body_size)
            run_title.font.color.rgb = primary_rgb
            
            run_comp = p_left.add_run(f" | {job.get('company', '')} ({job.get('location', '')})")
            run_comp.font.name = font_name
            run_comp.font.size = Pt(body_size)
            run_comp.font.color.rgb = RGBColor(71, 85, 105)
            
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(1)
            run_dates = p_right.add_run(job.get('dates', ''))
            run_dates.bold = True
            run_dates.font.name = font_name
            run_dates.font.size = Pt(body_size - 0.5)
            run_dates.font.color.rgb = RGBColor(71, 85, 105)
            
            set_cell_margins(cell_left, 0, 0, 0, 0)
            set_cell_margins(cell_right, 0, 0, 0, 0)
            
            # Bullets
            for bullet in job.get('bullet_points', []):
                bp = container.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.line_spacing = 1.1
                bp.paragraph_format.left_indent = Inches(0.20)
                
                parts = bullet.split('**')
                for idx, part in enumerate(parts):
                    run_part = bp.add_run(part)
                    run_part.font.name = font_name
                    run_part.font.size = Pt(body_size - 0.5)
                    if idx % 2 == 1:
                        run_part.bold = True

def render_projects_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, accent_style="underline", primary_color_hex="000000"):
    projects = data.get('projects', [])
    if projects:
        add_section_heading_docx(container, "Projects", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        for proj in projects:
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run_name = p.add_run(proj.get('name', ''))
            run_name.bold = True
            run_name.font.name = font_name
            run_name.font.size = Pt(body_size)
            run_name.font.color.rgb = primary_rgb
            
            if proj.get('role'):
                run_role = p.add_run(f" ({proj.get('role', '')})")
                run_role.italic = True
                run_role.font.name = font_name
                run_role.font.size = Pt(body_size - 0.5)
                
            for bullet in proj.get('description', []):
                bp = container.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.line_spacing = 1.1
                bp.paragraph_format.left_indent = Inches(0.20)
                
                parts = bullet.split('**')
                for idx, part in enumerate(parts):
                    run_part = bp.add_run(part)
                    run_part.font.name = font_name
                    run_part.font.size = Pt(body_size - 0.5)
                    if idx % 2 == 1:
                        run_part.bold = True

def render_education_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, width_col1=4.5, width_col2=2.0, accent_style="underline", primary_color_hex="000000"):
    edu_list = data.get('education', [])
    if edu_list:
        add_section_heading_docx(container, "Education", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        for edu in edu_list:
            table = container.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(width_col1)
            table.columns[1].width = Inches(width_col2)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            
            run_deg = p_left.add_run(edu.get('degree', ''))
            run_deg.bold = True
            run_deg.font.name = font_name
            run_deg.font.size = Pt(body_size)
            
            run_inst = p_left.add_run(f", {edu.get('institution', '')}")
            run_inst.font.name = font_name
            run_inst.font.size = Pt(body_size)
            
            if edu.get('gpa_or_honors'):
                run_gpa = p_left.add_run(f" ({edu.get('gpa_or_honors', '')})")
                run_gpa.font.name = font_name
                run_gpa.font.size = Pt(body_size - 0.5)
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            
            run_dates = p_right.add_run(edu.get('dates', ''))
            run_dates.font.name = font_name
            run_dates.font.size = Pt(body_size - 0.5)
            
            set_cell_margins(cell_left, 0, 0, 0, 0)
            set_cell_margins(cell_right, 0, 0, 0, 0)

def render_certifications_docx(container, data, font_name, primary_rgb, body_size=9.5, spacer_before=12, accent_style="underline", primary_color_hex="000000"):
    certs = data.get('certifications', [])
    if certs:
        add_section_heading_docx(container, "Certifications", font_name, primary_rgb, spacer_before, accent_style, primary_color_hex)
        p = container.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(", ".join(certs))
        run.font.name = font_name
        run.font.size = Pt(body_size)


def generate_docx(data, style_config=None):
    """
    Builds a beautifully styled Word Resume document matching the given style configuration.
    Returns: (bytes, adjustment_applied)
    """
    if style_config is None:
        style_config = TEMPLATES["modern_tech"]
        
    doc = Document()
    
    # Calculate word count for adaptive compression
    word_count = get_resume_word_count(data)
    adjustment_applied = False
    
    # Base layout variables
    layout = style_config.get("layout", "single_column")
    density = style_config.get("density", "spacious")
    accent_style = style_config.get("accent_style", "underline")
    profile_photo_bytes = style_config.get("profile_photo", None)
    
    # Check layout overflow constraints
    if layout == "two_column_sidebar" and word_count > 520:
        layout = "single_column" # Fallback to single column to avoid overflow
        adjustment_applied = True
        
    # Check manual density / adaptive sizing
    if density == "compact":
        body_size = 8.5
        spacer_before = 6
        m_val = 0.50
        if word_count > 600:
            body_size = 8.0
            spacer_before = 4
            adjustment_applied = True
    else: # spacious
        body_size = 10.0
        spacer_before = 12
        m_val = 0.80
        if word_count > 600:
            body_size = 9.0
            spacer_before = 8
            m_val = 0.65
            adjustment_applied = True
        elif word_count > 450:
            body_size = 9.5
            spacer_before = 10
            m_val = 0.70
            adjustment_applied = True
        
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(m_val)
        section.bottom_margin = Inches(m_val)
        section.left_margin = Inches(m_val)
        section.right_margin = Inches(m_val)
        
    font_name = style_config.get("font_docx", "Arial")
    primary_color_hex = style_config.get("primary_color", "#000000")
    primary_rgb = hex_to_rgb(primary_color_hex)
    
    # Check dark background contrast issue for docx
    bg_color_hex = style_config.get("bg_color", "#ffffff").lower()
    if bg_color_hex in ["#0f172a", "#1e293b", "#000000"]:
        # Dark theme in Word requires white/light text
        text_rgb = hex_to_rgb("#f8fafc")
    else:
        text_rgb = hex_to_rgb(style_config.get("text_color", "#334155"))
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = font_name
    style_normal.font.size = Pt(body_size)
    style_normal.font.color.rgb = text_rgb
    
    c_info = data.get('contact_info', {})
    
    # --- TWO COLUMN SIDEBAR LAYOUT ---
    if layout == "two_column_sidebar":
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.autofit = False
        
        # Printable width
        printable_w = 8.5 - (2 * m_val)
        sidebar_w = 2.1
        content_w = printable_w - sidebar_w
        
        layout_table.columns[0].width = Inches(sidebar_w)
        layout_table.columns[1].width = Inches(content_w)
        
        cell_sidebar = layout_table.cell(0, 0)
        cell_content = layout_table.cell(0, 1)
        
        set_cell_margins(cell_sidebar, top=0, bottom=0, left=0, right=144)
        set_cell_margins(cell_content, top=0, bottom=0, left=144, right=0)
        
        # Sidebar contact details
        p_side_c = cell_sidebar.paragraphs[0]
        
        # Profile Photo in sidebar if present
        if profile_photo_bytes:
            try:
                p_img = cell_sidebar.add_paragraph()
                p_img.paragraph_format.space_after = Pt(8)
                r_img = p_img.add_run()
                r_img.add_picture(io.BytesIO(profile_photo_bytes), width=Inches(1.2))
            except Exception:
                pass
                
        p_side_c.paragraph_format.space_after = Pt(2)
        r_name = p_side_c.add_run(c_info.get('full_name', 'Your Name'))
        r_name.font.size = Pt(16)
        r_name.font.bold = True
        r_name.font.color.rgb = primary_rgb
        
        for k in ['email', 'phone', 'location', 'linkedin', 'portfolio']:
            if c_info.get(k):
                p_item = cell_sidebar.add_paragraph()
                p_item.paragraph_format.space_after = Pt(2)
                run_i = p_item.add_run(c_info[k])
                run_i.font.size = Pt(body_size - 1.0)
                run_i.font.color.rgb = text_rgb
                
        # Sidebar modules
        render_skills_docx(cell_sidebar, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex)
        render_certifications_docx(cell_sidebar, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex)
        
        # Main content area
        p_title = cell_content.paragraphs[0]
        p_title.paragraph_format.space_after = Pt(8)
        run_name = p_title.add_run(c_info.get('full_name', 'Your Name').upper())
        run_name.font.size = Pt(20)
        run_name.font.bold = True
        run_name.font.color.rgb = primary_rgb
        
        sections_map = {
            "summary": lambda cont: render_summary_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex),
            "experience": lambda cont: render_experience_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, width_col1=(content_w-1.8), width_col2=1.8, accent_style=accent_style, primary_color_hex=primary_color_hex),
            "projects": lambda cont: render_projects_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex),
            "education": lambda cont: render_education_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, width_col1=(content_w-1.8), width_col2=1.8, accent_style=accent_style, primary_color_hex=primary_color_hex)
        }
        
        for section_id in style_config.get("section_order", []):
            if section_id in sections_map:
                sections_map[section_id](cell_content)
                
    # --- SINGLE COLUMN LAYOUT ---
    else:
        printable_w = 8.5 - (2 * m_val)
        w_l = printable_w - 1.8
        w_r = 1.8
        
        contact_parts = []
        for key in ['email', 'phone', 'location', 'linkedin', 'portfolio']:
            if c_info.get(key):
                contact_parts.append(c_info[key])
                
        # Header layout with profile photo if toggled
        if profile_photo_bytes:
            try:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(printable_w - 1.2)
                table.columns[1].width = Inches(1.2)
                
                cell_left = table.cell(0, 0)
                cell_right = table.cell(0, 1)
                
                set_cell_margins(cell_left, 0, 0, 0, 0)
                set_cell_margins(cell_right, 0, 0, 0, 0)
                
                # Name and contact on left
                name_p = cell_left.paragraphs[0]
                name_run = name_p.add_run(c_info.get('full_name', 'Your Name'))
                name_run.font.size = Pt(22)
                name_run.font.bold = True
                name_run.font.color.rgb = primary_rgb
                
                contact_p = cell_left.add_paragraph()
                contact_p.paragraph_format.space_after = Pt(12)
                contact_run = contact_p.add_run("  |  ".join(contact_parts))
                contact_run.font.size = Pt(body_size - 0.5)
                contact_run.font.color.rgb = text_rgb
                
                # Image on right
                p_img = cell_right.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_img = p_img.add_run()
                r_img.add_picture(io.BytesIO(profile_photo_bytes), width=Inches(1.0))
            except Exception:
                # Fallback to no photo
                profile_photo_bytes = None
                
        if not profile_photo_bytes:
            name_p = doc.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_p.paragraph_format.space_after = Pt(2)
            name_run = name_p.add_run(c_info.get('full_name', 'Your Name'))
            name_run.font.size = Pt(22)
            name_run.font.bold = True
            name_run.font.color.rgb = primary_rgb
            
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_p.paragraph_format.space_after = Pt(12)
            contact_run = contact_p.add_run("  |  ".join(contact_parts))
            contact_run.font.size = Pt(body_size - 0.5)
            contact_run.font.color.rgb = text_rgb
            
        sections_map = {
            "summary": lambda cont: render_summary_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex),
            "skills": lambda cont: render_skills_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex),
            "experience": lambda cont: render_experience_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, width_col1=w_l, width_col2=w_r, accent_style=accent_style, primary_color_hex=primary_color_hex),
            "projects": lambda cont: render_projects_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex),
            "education": lambda cont: render_education_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, width_col1=w_l, width_col2=w_r, accent_style=accent_style, primary_color_hex=primary_color_hex),
            "certifications": lambda cont: render_certifications_docx(cont, data, font_name, primary_rgb, body_size, spacer_before, accent_style, primary_color_hex)
        }
        
        for section_id in style_config.get("section_order", []):
            if section_id in sections_map:
                sections_map[section_id](doc)
                
    # Save output
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue(), adjustment_applied


# ==============================================================================
# PDF MODULAR BUILDING SYSTEM (REPORTLAB)
# ==============================================================================

def clean_markdown_bold(text):
    """
    Converts **bold** syntax to <b>bold</b> for ReportLab Paragraphs.
    """
    parts = text.split('**')
    for idx in range(len(parts)):
        if idx % 2 == 1:
            parts[idx] = f"<b>{parts[idx]}</b>"
    return "".join(parts)

def get_section_header_pdf(title, heading_style, primary_color_hex, accent_style="underline", spacer_before=8):
    p = Paragraph(title.upper(), heading_style)
    if accent_style == "minimal":
        return [p, Spacer(1, 4)]
    elif accent_style == "left_border":
        title_table = Table([[Spacer(1, 1), p]], colWidths=[4, None])
        title_table.setStyle(TableStyle([
            ('LINELEFT', (0, 0), (0, -1), 3.0, colors.HexColor(primary_color_hex)),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('LEFTPADDING', (0,0), (0,0), 0),
            ('RIGHTPADDING', (0,0), (0,0), 0),
            ('LEFTPADDING', (1,0), (1,0), 6),
            ('RIGHTPADDING', (1,0), (1,0), 0),
        ]))
        return [title_table, Spacer(1, 4)]
    else: # "underline"
        hr_color = colors.HexColor('#cbd5e1')
        return [
            p,
            HRFlowable(width="100%", thickness=0.8, color=hr_color, spaceBefore=1, spaceAfter=4)
        ]

def get_summary_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, accent_style="underline"):
    story = []
    summary = data.get('professional_summary', '')
    if summary:
        heading_style = ParagraphStyle('H_Sum', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Sum', parent=styles['Normal'], fontName=font_name, fontSize=body_size, leading=body_size+3, textColor=colors.HexColor(text_color_hex), spaceAfter=5)
        
        story.extend(get_section_header_pdf("Professional Summary", heading_style, primary_color_hex, accent_style, spacer_before))
        story.append(Paragraph(summary, body_style))
    return story

def get_skills_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, sidebar=False, accent_style="underline"):
    story = []
    skills_data = data.get('skills', {})
    if skills_data:
        heading_style = ParagraphStyle('H_Skl', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Skl', parent=styles['Normal'], fontName=font_name, fontSize=body_size - 0.5, leading=body_size+2.0, textColor=colors.HexColor(text_color_hex), spaceAfter=4)
        
        story.extend(get_section_header_pdf("Skills", heading_style, primary_color_hex, accent_style, spacer_before))
        
        tech_skills = skills_data.get('technical_skills', [])
        soft_skills = skills_data.get('soft_skills', [])
        tools = skills_data.get('tools_frameworks', [])
        
        if sidebar:
            if tech_skills:
                story.append(Paragraph("<b>Technical Skills:</b>", body_style))
                story.append(Paragraph(", ".join(tech_skills), ParagraphStyle('T_S', parent=body_style, leftIndent=5, spaceAfter=4)))
            if soft_skills:
                story.append(Paragraph("<b>Soft Skills:</b>", body_style))
                story.append(Paragraph(", ".join(soft_skills), ParagraphStyle('T_S2', parent=body_style, leftIndent=5, spaceAfter=4)))
            if tools:
                story.append(Paragraph("<b>Tools:</b>", body_style))
                story.append(Paragraph(", ".join(tools), ParagraphStyle('T_S3', parent=body_style, leftIndent=5, spaceAfter=4)))
        else:
            skills_text = ""
            if tech_skills:
                skills_text += f"<b>Technical Skills:</b> {', '.join(tech_skills)}<br/>"
            if soft_skills:
                skills_text += f"<b>Soft Skills:</b> {', '.join(soft_skills)}<br/>"
            if tools:
                skills_text += f"<b>Tools & Systems:</b> {', '.join(tools)}"
            story.append(Paragraph(skills_text, body_style))
            
    return story

def get_experience_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, width_left=360, width_right=160, accent_style="underline"):
    story = []
    jobs = data.get('work_experience', [])
    if jobs:
        heading_style = ParagraphStyle('H_Exp', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Exp', parent=styles['Normal'], fontName=font_name, fontSize=body_size, leading=body_size+3, textColor=colors.HexColor(text_color_hex), spaceAfter=4)
        bullet_style = ParagraphStyle('Bul_Exp', parent=body_style, fontSize=body_size - 0.5, leading=body_size+2.0, leftIndent=12, firstLineIndent=-8, spaceAfter=2)
        
        story.extend(get_section_header_pdf("Professional Experience", heading_style, primary_color_hex, accent_style, spacer_before))
        for job in jobs:
            title_text = f"<b>{job.get('job_title', '')}</b> | {job.get('company', '')} ({job.get('location', '')})"
            date_text = f"<b>{job.get('dates', '')}</b>"
            
            job_p_left = Paragraph(title_text, ParagraphStyle('JobL', parent=body_style, textColor=colors.HexColor(primary_color_hex)))
            job_p_right = Paragraph(date_text, ParagraphStyle('JobR', parent=body_style, alignment=TA_CENTER, textColor=colors.HexColor(text_color_hex)))
            
            t = Table([[job_p_left, job_p_right]], colWidths=[width_left, width_right])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t)
            story.append(Spacer(1, 2))
            
            for bullet in job.get('bullet_points', []):
                cleaned_bullet = clean_markdown_bold(bullet)
                story.append(Paragraph(f"&bull; {cleaned_bullet}", bullet_style))
                
            story.append(Spacer(1, 4))
            
    return story

def get_projects_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, accent_style="underline"):
    story = []
    projects = data.get('projects', [])
    if projects:
        heading_style = ParagraphStyle('H_Prj', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Prj', parent=styles['Normal'], fontName=font_name, fontSize=body_size, leading=body_size+3, textColor=colors.HexColor(text_color_hex), spaceAfter=4)
        bullet_style = ParagraphStyle('Bul_Prj', parent=body_style, fontSize=body_size - 0.5, leading=body_size+2.0, leftIndent=12, firstLineIndent=-8, spaceAfter=2)
        
        story.extend(get_section_header_pdf("Projects", heading_style, primary_color_hex, accent_style, spacer_before))
        for project in projects:
            role_suffix = f" - <i>{project.get('role', '')}</i>" if project.get('role') else ""
            story.append(Paragraph(f"<b>{project.get('name', '')}</b>{role_suffix}", ParagraphStyle('ProjT', parent=body_style, textColor=colors.HexColor(primary_color_hex))))
            
            for bullet in project.get('description', []):
                cleaned_bullet = clean_markdown_bold(bullet)
                story.append(Paragraph(f"&bull; {cleaned_bullet}", bullet_style))
            story.append(Spacer(1, 4))
            
    return story

def get_education_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, width_left=380, width_right=140, accent_style="underline"):
    story = []
    edu_list = data.get('education', [])
    if edu_list:
        heading_style = ParagraphStyle('H_Edu', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Edu', parent=styles['Normal'], fontName=font_name, fontSize=body_size, leading=body_size+3, textColor=colors.HexColor(text_color_hex), spaceAfter=4)
        
        story.extend(get_section_header_pdf("Education", heading_style, primary_color_hex, accent_style, spacer_before))
        for edu in edu_list:
            edu_details = f"<b>{edu.get('degree', '')}</b>, {edu.get('institution', '')}"
            if edu.get('gpa_or_honors'):
                edu_details += f" ({edu.get('gpa_or_honors', '')})"
            
            edu_p_left = Paragraph(edu_details, body_style)
            edu_p_right = Paragraph(f"<b>{edu.get('dates', '')}</b>", ParagraphStyle('EduR', parent=body_style, alignment=TA_CENTER, textColor=colors.HexColor(text_color_hex)))
            
            t = Table([[edu_p_left, edu_p_right]], colWidths=[width_left, width_right])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t)
            story.append(Spacer(1, 2))
        story.append(Spacer(1, 4))
        
    return story

def get_certifications_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size=9.5, spacer_before=8, accent_style="underline"):
    story = []
    certs = data.get('certifications', [])
    if certs:
        heading_style = ParagraphStyle('H_Crt', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=spacer_before, spaceAfter=2, keepWithNext=True)
        body_style = ParagraphStyle('B_Crt', parent=styles['Normal'], fontName=font_name, fontSize=body_size, leading=body_size+3, textColor=colors.HexColor(text_color_hex), spaceAfter=4)
        
        story.extend(get_section_header_pdf("Certifications", heading_style, primary_color_hex, accent_style, spacer_before))
        story.append(Paragraph(", ".join(certs), body_style))
        
    return story


def generate_pdf(data, style_config=None):
    """
    Creates an executive, clean resume PDF from tailored AI JSON data.
    Returns: (bytes, adjustment_applied)
    """
    try:
        from reportlab.platypus import Image
    except ImportError:
        Image = None

    if style_config is None:
        style_config = TEMPLATES["modern_tech"]
        
    pdf_io = io.BytesIO()
    
    # Calculate word count for adaptive compression
    word_count = get_resume_word_count(data)
    adjustment_applied = False
    
    # Base layout variables
    layout = style_config.get("layout", "single_column")
    density = style_config.get("density", "spacious")
    accent_style = style_config.get("accent_style", "underline")
    profile_photo_bytes = style_config.get("profile_photo", None)
    
    # Check layout overflow constraints
    if layout == "two_column_sidebar" and word_count > 520:
        layout = "single_column" # Fallback to single column to avoid overflow
        adjustment_applied = True
        
    # Check manual density / adaptive sizing
    if density == "compact":
        body_size = 8.5
        spacer_before = 6
        m_val = 0.50
        if word_count > 600:
            body_size = 7.8
            spacer_before = 4
            adjustment_applied = True
    else: # spacious
        body_size = 10.0
        spacer_before = 12
        m_val = 0.80
        if word_count > 600:
            body_size = 9.0
            spacer_before = 8
            m_val = 0.65
            adjustment_applied = True
        elif word_count > 450:
            body_size = 9.5
            spacer_before = 10
            m_val = 0.70
            adjustment_applied = True
            
    doc = SimpleDocTemplate(
        pdf_io, 
        pagesize=letter, 
        leftMargin=m_val * 72, 
        rightMargin=m_val * 72, 
        topMargin=m_val * 72, 
        bottomMargin=m_val * 72
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    # Configure font mapping (ensure only safe PDF core fonts compile)
    font_name = style_config.get("font_pdf", "Helvetica")
    if font_name not in ["Helvetica", "Times-Roman", "Courier"]:
        font_name = "Helvetica"
        
    primary_color_hex = style_config.get("primary_color", "#000000")
    text_color_hex = style_config.get("text_color", "#334155")
    
    c_info = data.get('contact_info', {})
    
    # Text styles
    title_style = ParagraphStyle(
        'T_Name',
        parent=styles['Normal'],
        fontName=f"{font_name}-Bold",
        fontSize=18 if body_size < 9.0 else 20,
        leading=22 if body_size < 9.0 else 24,
        textColor=colors.HexColor(primary_color_hex),
        spaceAfter=3
    )
    
    subtitle_style = ParagraphStyle(
        'T_Sub',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=body_size - 0.5,
        leading=body_size + 2.0,
        textColor=colors.HexColor(text_color_hex),
        spaceAfter=8
    )

    # Load Profile Image if present
    img = None
    if profile_photo_bytes and Image:
        try:
            img = Image(io.BytesIO(profile_photo_bytes), width=72, height=72)
        except Exception:
            pass

    try:
        # --- TWO COLUMN SIDEBAR LAYOUT ---
        if layout == "two_column_sidebar":
            total_w = 612 - (2 * m_val * 72)
            sidebar_w = 165
            content_w = total_w - sidebar_w
            
            left_story = []
            c_title_style = ParagraphStyle('SideC_T', parent=styles['Normal'], fontName=f"{font_name}-Bold", fontSize=11, leading=14, textColor=colors.HexColor(primary_color_hex), spaceBefore=4, spaceAfter=2)
            c_body_style = ParagraphStyle('SideC_B', parent=styles['Normal'], fontName=font_name, fontSize=body_size - 1.0, leading=body_size+1.5, textColor=colors.HexColor(text_color_hex), spaceAfter=2)
            
            # Profile Photo in sidebar if present
            if img:
                left_story.append(img)
                left_story.append(Spacer(1, 8))
                
            left_story.extend(get_section_header_pdf("Contact", c_title_style, primary_color_hex, accent_style, spacer_before))
            for k_id, k_lbl in [('email', 'Email'), ('phone', 'Phone'), ('location', 'Location'), ('linkedin', 'LinkedIn'), ('portfolio', 'Portfolio')]:
                if c_info.get(k_id):
                    left_story.append(Paragraph(f"<b>{k_lbl}:</b> {c_info[k_id]}", c_body_style))
                    
            left_story.append(Spacer(1, 6))
            left_story.extend(get_skills_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, sidebar=True, accent_style=accent_style))
            left_story.append(Spacer(1, 6))
            left_story.extend(get_certifications_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style))
            
            right_story = []
            right_story.append(Paragraph(c_info.get('full_name', 'Your Name'), title_style))
            right_story.append(Spacer(1, 4))
            
            sections_map = {
                "summary": lambda: get_summary_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style),
                "experience": lambda: get_experience_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, width_left=(content_w - 110), width_right=110, accent_style=accent_style),
                "projects": lambda: get_projects_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style),
                "education": lambda: get_education_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, width_left=(content_w - 100), width_right=100, accent_style=accent_style)
            }
            
            for section_id in style_config.get("section_order", []):
                if section_id in sections_map:
                    right_story.extend(sections_map[section_id]())
                    
            table_data = [[left_story, right_story]]
            t = Table(table_data, colWidths=[sidebar_w, content_w])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (0,0), 12),
                ('LEFTPADDING', (1,0), (1,0), 6),
                ('RIGHTPADDING', (1,0), (1,0), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t)
            
        # --- SINGLE COLUMN LAYOUT ---
        else:
            total_w = 612 - (2 * m_val * 72)
            w_l = total_w - 110
            w_r = 110
            
            contact_parts = []
            for key in ['email', 'phone', 'location', 'linkedin', 'portfolio']:
                if c_info.get(key):
                    contact_parts.append(c_info[key])
                    
            # Profile Photo layout for single column
            if img:
                header_p1 = Paragraph(c_info.get('full_name', 'Your Name'), title_style)
                header_p2 = Paragraph("  |  ".join(contact_parts), ParagraphStyle('T_Sub_C_L', parent=subtitle_style, alignment=TA_LEFT))
                
                t_header = Table([[ [header_p1, header_p2], img ]], colWidths=[total_w - 80, 80])
                t_header.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0),
                ]))
                story.append(t_header)
                story.append(Spacer(1, 4))
            else:
                story.append(Paragraph(c_info.get('full_name', 'Your Name'), title_style))
                story.append(Paragraph("  |  ".join(contact_parts), ParagraphStyle('T_Sub_C', parent=subtitle_style, alignment=TA_CENTER)))
                story.append(Spacer(1, 4))
            
            sections_map = {
                "summary": lambda: get_summary_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style),
                "skills": lambda: get_skills_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, sidebar=False, accent_style=accent_style),
                "experience": lambda: get_experience_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, width_left=w_l, width_right=w_r, accent_style=accent_style),
                "projects": lambda: get_projects_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style),
                "education": lambda: get_education_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, width_left=(total_w - 100), width_right=100, accent_style=accent_style),
                "certifications": lambda: get_certifications_pdf(data, styles, font_name, primary_color_hex, text_color_hex, body_size, spacer_before, accent_style=accent_style)
            }
            
            for section_id in style_config.get("section_order", []):
                if section_id in sections_map:
                    story.extend(sections_map[section_id]())
                    
        # Build PDF doc
        doc.build(story)
        
    except Exception as e:
        # Emergency recovery: Build a very simple page that never fails
        story = []
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(pdf_io, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        
        # Simple name header
        story.append(Paragraph(c_info.get('full_name', 'Your Name'), ParagraphStyle('FallbackTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=18, leading=22)))
        
        contact_parts = [c_info[k] for k in ['email', 'phone', 'location'] if c_info.get(k)]
        story.append(Paragraph(" | ".join(contact_parts), ParagraphStyle('FallbackContact', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12)))
        story.append(Spacer(1, 10))
        
        # Add summary
        if data.get('professional_summary'):
            story.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", ParagraphStyle('FallbackSection', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10, leading=14)))
            story.append(Paragraph(data.get('professional_summary'), ParagraphStyle('FallbackBody', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12)))
            story.append(Spacer(1, 10))
            
        # Add experience
        if data.get('work_experience'):
            story.append(Paragraph("<b>PROFESSIONAL EXPERIENCE</b>", ParagraphStyle('FallbackSection', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10, leading=14)))
            for job in data.get('work_experience', []):
                story.append(Paragraph(f"<b>{job.get('job_title', '')}</b> at {job.get('company', '')}", ParagraphStyle('FallbackJob', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9, leading=12)))
                for bp in job.get('bullet_points', []):
                    story.append(Paragraph(f"&bull; {clean_markdown_bold(bp)}", ParagraphStyle('FallbackBP', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11, leftIndent=10)))
            story.append(Spacer(1, 10))
            
        doc.build(story)
        adjustment_applied = True
        
    pdf_io.seek(0)
    return pdf_io.getvalue(), adjustment_applied


# ==============================================================================
# ZIP COMPRESSION EXPORTER
# ==============================================================================

def generate_all_zip(tailored_data):
    """
    Creates a ZIP archive containing all templates in both PDF and DOCX formats.
    """
    zip_io = io.BytesIO()
    with zipfile.ZipFile(zip_io, 'w') as zf:
        for tid, t_config in TEMPLATES.items():
            safe_name = t_config["name"].replace(" ", "_")
            
            # 1. DOCX format
            try:
                docx_bytes, _ = generate_docx(tailored_data, t_config)
                zf.writestr(f"{safe_name}.docx", docx_bytes)
            except Exception as e:
                print(f"Failed packaging DOCX template '{tid}': {e}")
                
            # 2. PDF format
            try:
                pdf_bytes, _ = generate_pdf(tailored_data, t_config)
                zf.writestr(f"{safe_name}.pdf", pdf_bytes)
            except Exception as e:
                print(f"Failed packaging PDF template '{tid}': {e}")
                
    zip_io.seek(0)
    return zip_io.getvalue()
